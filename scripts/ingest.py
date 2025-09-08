#!/usr/bin/env python3
"""
Document Ingestion Pipeline
Processes documents from various formats into vector database
"""

import os
import sys
import json
import hashlib
import argparse
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
from dataclasses import dataclass

# Add parent directory to path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from vector_db import VectorDatabase
from chunking_strategy import DocumentChunker, ChunkConfig, ChunkingStrategy

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@dataclass
class IngestConfig:
    """Configuration for document ingestion"""
    input_dir: str = "./data/documents"
    output_dir: str = "./data/processed"
    file_types: List[str] = None
    batch_size: int = 10
    dedupe: bool = True
    chunk_size: int = 512
    chunk_overlap: int = 50
    max_file_size_mb: int = 100
    skip_errors: bool = True

class DocumentParser:
    """Multi-format document parser"""
    
    @staticmethod
    def parse_txt(file_path: str) -> Dict[str, Any]:
        """Parse plain text file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return {
            'content': content,
            'metadata': {
                'format': 'txt',
                'file_path': file_path,
                'file_size': os.path.getsize(file_path)
            }
        }
    
    @staticmethod
    def parse_pdf(file_path: str) -> Dict[str, Any]:
        """Parse PDF file using PyPDF2"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                content = ""
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        text = page.extract_text()
                        if text:
                            content += f"\n--- Page {page_num + 1} ---\n{text}"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                
                metadata = {
                    'format': 'pdf',
                    'file_path': file_path,
                    'file_size': os.path.getsize(file_path),
                    'page_count': len(reader.pages),
                    'title': reader.metadata.title if reader.metadata and reader.metadata.title else None,
                    'author': reader.metadata.author if reader.metadata and reader.metadata.author else None
                }
                
                return {'content': content.strip(), 'metadata': metadata}
                
        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2>=3.0.0")
            return {'content': '', 'metadata': {'format': 'pdf', 'file_path': file_path, 'error': 'PyPDF2 not installed'}}
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            return {'content': '', 'metadata': {'format': 'pdf', 'file_path': file_path, 'error': str(e)}}
    
    @staticmethod
    def parse_docx(file_path: str) -> Dict[str, Any]:
        """Parse DOCX file using python-docx"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            content = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content += text + "\n\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        content += row_text + "\n"
                content += "\n"
            
            # Extract metadata
            properties = doc.core_properties
            metadata = {
                'format': 'docx',
                'file_path': file_path,
                'file_size': os.path.getsize(file_path),
                'title': properties.title,
                'author': properties.author,
                'created': properties.created.isoformat() if properties.created else None,
                'modified': properties.modified.isoformat() if properties.modified else None,
                'subject': properties.subject,
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
            
            return {'content': content.strip(), 'metadata': metadata}
            
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx>=0.8.11")
            return {'content': '', 'metadata': {'format': 'docx', 'file_path': file_path, 'error': 'python-docx not installed'}}
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            return {'content': '', 'metadata': {'format': 'docx', 'file_path': file_path, 'error': str(e)}}
    
    @staticmethod
    def parse_md(file_path: str) -> Dict[str, Any]:
        """Parse Markdown file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return {
            'content': content,
            'metadata': {
                'format': 'md',
                'file_path': file_path,
                'file_size': os.path.getsize(file_path)
            }
        }
    
    @staticmethod
    def parse_html(file_path: str) -> Dict[str, Any]:
        """Parse HTML file using BeautifulSoup"""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract text content
            content = soup.get_text()
            
            # Extract metadata
            title = None
            if soup.title:
                title = soup.title.string
            
            # Look for meta tags
            description = None
            author = None
            meta_tags = soup.find_all('meta')
            for meta in meta_tags:
                if meta.get('name') == 'description':
                    description = meta.get('content')
                elif meta.get('name') == 'author':
                    author = meta.get('content')
            
            metadata = {
                'format': 'html',
                'file_path': file_path,
                'file_size': os.path.getsize(file_path),
                'title': title,
                'description': description,
                'author': author
            }
            
            return {'content': content.strip(), 'metadata': metadata}
            
        except ImportError:
            logger.error("BeautifulSoup not installed. Install with: pip install beautifulsoup4>=4.11.0")
            return {'content': '', 'metadata': {'format': 'html', 'file_path': file_path, 'error': 'beautifulsoup4 not installed'}}
        except Exception as e:
            logger.error(f"Error parsing HTML {file_path}: {e}")
            return {'content': '', 'metadata': {'format': 'html', 'file_path': file_path, 'error': str(e)}}
    
    @staticmethod
    def parse_csv(file_path: str) -> Dict[str, Any]:
        """Parse CSV file using pandas"""
        try:
            import pandas as pd
            
            df = pd.read_csv(file_path)
            
            # Convert to text representation
            content = f"CSV Data from {Path(file_path).name}:\n\n"
            content += f"Columns: {', '.join(df.columns.tolist())}\n\n"
            content += df.to_string(index=False)
            
            metadata = {
                'format': 'csv',
                'file_path': file_path,
                'file_size': os.path.getsize(file_path),
                'row_count': len(df),
                'column_count': len(df.columns),
                'columns': df.columns.tolist()
            }
            
            return {'content': content, 'metadata': metadata}
            
        except ImportError:
            logger.error("pandas not installed. Install with: pip install pandas>=2.0.3")
            return {'content': '', 'metadata': {'format': 'csv', 'file_path': file_path, 'error': 'pandas not installed'}}
        except Exception as e:
            logger.error(f"Error parsing CSV {file_path}: {e}")
            return {'content': '', 'metadata': {'format': 'csv', 'file_path': file_path, 'error': str(e)}}

class DocumentIngester:
    """Main ingestion pipeline"""
    
    def __init__(self, config: IngestConfig):
        self.config = config
        self.vector_db = VectorDatabase()
        self.chunker = DocumentChunker(
            ChunkConfig(
                chunk_size=config.chunk_size,
                overlap=config.chunk_overlap,
                strategy=ChunkingStrategy.HYBRID
            )
        )
        self.parser = DocumentParser()
        self.processed_hashes = set()
        self._load_processed_hashes()
    
    def _load_processed_hashes(self):
        """Load previously processed document hashes"""
        hash_file = Path(self.config.output_dir) / 'processed_hashes.json'
        if hash_file.exists():
            with open(hash_file, 'r') as f:
                self.processed_hashes = set(json.load(f))
            logger.info(f"Loaded {len(self.processed_hashes)} processed hashes")
    
    def _save_processed_hashes(self):
        """Save processed document hashes"""
        hash_file = Path(self.config.output_dir) / 'processed_hashes.json'
        os.makedirs(self.config.output_dir, exist_ok=True)
        with open(hash_file, 'w') as f:
            json.dump(list(self.processed_hashes), f)
    
    def _compute_hash(self, content: str) -> str:
        """Compute content hash for deduplication"""
        return hashlib.sha256(content.encode()).hexdigest()
    
    def _parse_document(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Parse document based on file type"""
        ext = Path(file_path).suffix.lower()
        
        parsers = {
            '.txt': self.parser.parse_txt,
            '.md': self.parser.parse_md,
            '.pdf': self.parser.parse_pdf,
            '.docx': self.parser.parse_docx,
            '.html': self.parser.parse_html,
            '.htm': self.parser.parse_html,
            '.csv': self.parser.parse_csv,
        }
        
        parser = parsers.get(ext)
        if not parser:
            logger.warning(f"No parser for file type: {ext}")
            return None
        
        try:
            return parser(file_path)
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            if not self.config.skip_errors:
                raise
            return None
    
    def ingest_file(self, file_path: str) -> bool:
        """Ingest a single file"""
        logger.info(f"Processing: {file_path}")
        
        # Parse document
        doc_data = self._parse_document(file_path)
        if not doc_data or not doc_data.get('content'):
            return False
        
        # Check deduplication
        content_hash = self._compute_hash(doc_data['content'])
        if self.config.dedupe and content_hash in self.processed_hashes:
            logger.info(f"Skipping duplicate: {file_path}")
            return False
        
        # Chunk document using advanced chunking strategy
        try:
            chunks = self.chunker.chunk_document(doc_data['content'], doc_data['metadata'])
            logger.info(f"Document chunked into {len(chunks)} chunks")
        except Exception as e:
            logger.error(f"Chunking failed: {e}")
            # Fallback to simple chunking
            content = doc_data['content']
            chunks = []
            for i in range(0, len(content), self.config.chunk_size):
                chunk_content = content[i:i + self.config.chunk_size]
                if chunk_content.strip():
                    from chunking_strategy import DocumentChunk
                    chunks.append(DocumentChunk(
                        content=chunk_content,
                        metadata=doc_data['metadata'],
                        chunk_index=len(chunks),
                        token_count=len(chunk_content) // 4,  # Rough estimate
                        char_start=i,
                        char_end=i + len(chunk_content)
                    ))
        
        # Add to vector database
        for i, chunk in enumerate(chunks):
            # Handle both DocumentChunk objects and simple strings
            if hasattr(chunk, 'content'):
                chunk_content = chunk.content
                chunk_metadata = chunk.metadata.copy()
                chunk_metadata.update({
                    'chunk_index': chunk.chunk_index,
                    'token_count': chunk.token_count,
                    'char_start': chunk.char_start,
                    'char_end': chunk.char_end
                })
            else:
                chunk_content = chunk
                chunk_metadata = doc_data['metadata'].copy()
                chunk_metadata.update({
                    'chunk_index': i,
                    'token_count': len(chunk) // 4
                })
            
            # Add title and source if not present
            chunk_title = chunk_metadata.get('title') or Path(file_path).stem
            if chunk_metadata.get('chunk_index', 0) > 0:
                chunk_title += f" (chunk {chunk_metadata['chunk_index'] + 1})"
            
            doc_id = self.vector_db.add_document(
                content=chunk_content,
                title=chunk_title,
                source=file_path
            )
            logger.info(f"Added chunk {i+1}/{len(chunks)}, ID: {doc_id}")
        
        # Mark as processed
        self.processed_hashes.add(content_hash)
        return True
    
    def ingest_directory(self, directory: str) -> Dict[str, Any]:
        """Ingest all documents in a directory"""
        results = {
            'processed': 0,
            'skipped': 0,
            'errors': 0,
            'files': []
        }
        
        # Get file list
        file_types = self.config.file_types or ['txt', 'md', 'pdf', 'docx', 'html', 'htm', 'csv']
        patterns = [f"*.{ft}" for ft in file_types]
        
        files = []
        for pattern in patterns:
            files.extend(Path(directory).rglob(pattern))
        
        logger.info(f"Found {len(files)} files to process")
        
        # Process in batches
        for i in range(0, len(files), self.config.batch_size):
            batch = files[i:i+self.config.batch_size]
            logger.info(f"Processing batch {i//self.config.batch_size + 1}")
            
            for file_path in batch:
                try:
                    if self.ingest_file(str(file_path)):
                        results['processed'] += 1
                    else:
                        results['skipped'] += 1
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {e}")
                    results['errors'] += 1
                
                results['files'].append(str(file_path))
            
            # Save progress after each batch
            self._save_processed_hashes()
        
        return results

def main():
    """Main entry point"""
    parser = argparse.ArgumentParser(description='Document Ingestion Pipeline')
    parser.add_argument('--input-dir', default='./data/documents', help='Input directory')
    parser.add_argument('--output-dir', default='./data/processed', help='Output directory')
    parser.add_argument('--file-types', nargs='+', help='File types to process')
    parser.add_argument('--batch-size', type=int, default=10, help='Batch size')
    parser.add_argument('--chunk-size', type=int, default=512, help='Chunk size')
    parser.add_argument('--chunk-overlap', type=int, default=50, help='Chunk overlap')
    parser.add_argument('--dedupe', action='store_true', help='Enable deduplication')
    parser.add_argument('--file', help='Process single file')
    
    args = parser.parse_args()
    
    # Create config
    config = IngestConfig(
        input_dir=args.input_dir,
        output_dir=args.output_dir,
        file_types=args.file_types,
        batch_size=args.batch_size,
        chunk_size=args.chunk_size,
        chunk_overlap=args.chunk_overlap,
        dedupe=args.dedupe
    )
    
    # Initialize ingester
    ingester = DocumentIngester(config)
    
    # Process
    if args.file:
        success = ingester.ingest_file(args.file)
        print(f"Ingestion {'successful' if success else 'failed'}")
    else:
        results = ingester.ingest_directory(args.input_dir)
        print(f"\nIngestion Results:")
        print(f"  Processed: {results['processed']}")
        print(f"  Skipped: {results['skipped']}")
        print(f"  Errors: {results['errors']}")
        
        # Save results
        results_file = Path(config.output_dir) / f"ingest_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(results_file, 'w') as f:
            json.dump(results, f, indent=2)
        print(f"\nResults saved to: {results_file}")

if __name__ == '__main__':
    main()
